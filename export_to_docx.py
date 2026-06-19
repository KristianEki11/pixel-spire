import os
import re
from docx import Document

frontend_files = [
    'index.html',
    'style.css',
    'ui.js',
]

backend_files = [
    'engine.js',
    'game.js',
    'data/cards.js',
    'data/curses.js',
    'data/enemies.js',
    'data/events.js',
    'data/relics.js',
    'data/stages.js'
]

def create_docx(title, files, output_filename):
    document = Document()
    document.add_heading(title, 0)

    for filepath in files:
        document.add_heading(filepath, level=1)
        
        if os.path.exists(filepath):
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # Split content into blocks by one or more blank lines
                blocks = re.split(r'\n\s*\n', content)
                
                for block in blocks:
                    block = block.strip()
                    if not block:
                        continue
                    
                    # Add a 1x1 table for each block
                    table = document.add_table(rows=1, cols=1)
                    table.style = 'Table Grid'
                    cell = table.cell(0, 0)
                    
                    # Add code block with default text color
                    p = cell.paragraphs[0]
                    run = p.add_run(block)
                    run.font.name = 'Courier New'
                    
                    # Add some space after the table
                    document.add_paragraph()
                
                document.add_page_break()
            except Exception as e:
                document.add_paragraph(f"Error reading file: {e}")
        else:
            document.add_paragraph(f"File not found: {filepath}")

    document.save(output_filename)
    print(f"Successfully generated {output_filename}")

create_docx('Frontend Source Code - Pixel Spire', frontend_files, 'SourceCode_PixelSpire_Frontend.docx')
create_docx('Backend Source Code - Pixel Spire', backend_files, 'SourceCode_PixelSpire_Backend.docx')


